from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse
from docx import Document
import fitz  # PyMuPDF for PDFs
import openai
import os
import tempfile

app = FastAPI()

# Set your OpenAI API key
openai.api_key = "sk-...gDMA"

@app.post("/process-file/")
async def process_file(file: UploadFile = File(...)):
    # Check file type
    if file.content_type not in ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "text/plain"]:
        return {"error": "Unsupported file type. Please upload a PDF, DOCX, or TXT file."}

    # Extract text based on file type
    content = await file.read()
    if file.content_type == "application/pdf":
        text = extract_text_from_pdf(content)
    elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text = extract_text_from_docx(content)
    elif file.content_type == "text/plain":
        text = content.decode("utf-8")
    else:
        return {"error": "Unsupported file type."}

    # Process the text with OpenAI
    chronology = generate_chronology(text)

    # Generate Word document
    word_path = create_word_doc(chronology)

    # Return the Word file
    return FileResponse(word_path, filename="chronology.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def extract_text_from_pdf(content):
    """Extract text from a PDF file."""
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    temp_file.write(content)
    temp_file.close()

    text = ""
    pdf_doc = fitz.open(temp_file.name)
    for page in pdf_doc:
        text += page.get_text()
    pdf_doc.close()
    os.unlink(temp_file.name)  # Clean up temporary file
    return text


def extract_text_from_docx(content):
    """Extract text from a DOCX file."""
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    temp_file.write(content)
    temp_file.close()

    doc = Document(temp_file.name)
    os.unlink(temp_file.name)  # Clean up temporary file
    return " ".join([p.text for p in doc.paragraphs])


def generate_chronology(text):
    """Use OpenAI to extract a chronology."""
    prompt = f"""
    Extract a chronology from the following text with columns: Date, Event, and Analysis.
    {text}
    """

    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}]
    )
    return eval(response.choices[0].message.content)  # Assuming JSON-like output


def create_word_doc(chronology):
    """Generate a Word document with the chronology."""
    doc = Document()
    doc.add_heading("Chronology", level=1)

    # Add a table
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Date"
    hdr_cells[1].text = "Event"
    hdr_cells[2].text = "Analysis"

    # Populate the table
    for row in chronology:
        cells = table.add_row().cells
        cells[0].text = row.get("Date", "N/A")
        cells[1].text = row.get("Event", "N/A")
        cells[2].text = row.get("Analysis", "N/A")

    # Save the document to a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name

