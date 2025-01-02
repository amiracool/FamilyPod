from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse
from docx import Document
#import exceptions
import openai
import os
import tempfile

app = FastAPI()

# Set your OpenAI API key
openai.api_key = "sk-...gDMA"

@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    # Extract text from the file
    text = (await file.read()).decode("utf-8")

    # Process the text with OpenAI
    chronology = generate_chronology(text)

    # Create a Word document
    word_file_path = create_word_document(chronology)

    # Serve the Word file for download
    return FileResponse(word_file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="chronology.docx")

def generate_chronology(text):
    # Prompt OpenAI to generate chronology
    prompt = f"""
    Extract a chronology from this text with columns Date, Event, and Analysis.

    Text:
    {text}

    Return as JSON.
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "You are an assistant that organizes text into chronologies."},
                  {"role": "user", "content": prompt}],
    )
    return eval(response["choices"][0]["message"]["content"])

def create_word_document(chronology):
    # Create a Word document
    doc = Document()
    doc.add_heading("Chronology", level=1)

    # Add a table
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Date"
    hdr_cells[1].text = "Event"
    hdr_cells[2].text = "Analysis"

    # Populate the table
    for entry in chronology:
        row_cells = table.add_row().cells
        row_cells[0].text = entry.get("Date", "N/A")
        row_cells[1].text = entry.get("Event", "N/A")
        row_cells[2].text = entry.get("Analysis", "N/A")

    # Save the file to a temporary location
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, "chronology.docx")
    doc.save(file_path)
    return file_path
